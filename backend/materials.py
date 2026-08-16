"""User-supplied research materials: PDFs, Word documents, plain text, HTML.

Files are uploaded before the debate starts, parsed to plain text immediately
(so problems surface at upload time, not mid-debate) and staged in memory.
When a debate starts they become research Docs, indexed for retrieval exactly
like the web sources.
"""
from __future__ import annotations

import io
import threading
import uuid
from dataclasses import dataclass

from .research import Doc

MAX_FILES = 12
MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_MATERIAL_CHARS = 200_000

TEXT_EXTENSIONS = {"txt", "md", "markdown", "rst", "csv", "log"}
ACCEPTED_EXTENSIONS = TEXT_EXTENSIONS | {"pdf", "docx", "html", "htm"}


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx_text(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("  ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _html_text(data: bytes) -> str:
    import trafilatura
    html = data.decode("utf-8", errors="replace")
    return trafilatura.extract(html) or ""


def extract_text(filename: str, data: bytes) -> str:
    """Parse an uploaded file to plain text. Raises ValueError with a
    user-facing message on anything unusable."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in ACCEPTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '.{ext}'. Accepted: "
            + ", ".join("." + e for e in sorted(ACCEPTED_EXTENSIONS)))
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(
            f"File is too large ({len(data) / 1024**2:.1f} MB, "
            f"max {MAX_FILE_BYTES // 1024**2} MB).")
    try:
        if ext == "pdf":
            text = _pdf_text(data)
        elif ext == "docx":
            text = _docx_text(data)
        elif ext in ("html", "htm"):
            text = _html_text(data)
        else:
            text = data.decode("utf-8", errors="replace")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Could not read '{filename}': {e}")
    text = text.strip()
    if len(text) < 100:
        raise ValueError(
            f"'{filename}' contains almost no extractable text. "
            "Scanned/image-only PDFs are not supported.")
    return text[:MAX_MATERIAL_CHARS]


@dataclass
class Material:
    id: str
    filename: str
    text: str


class MaterialStore:
    """In-memory staging area for uploaded materials (cleared on restart)."""

    def __init__(self):
        self._items: dict[str, Material] = {}
        self._lock = threading.Lock()

    def add(self, filename: str, data: bytes) -> Material:
        with self._lock:
            if len(self._items) >= MAX_FILES:
                raise ValueError(f"At most {MAX_FILES} files can be staged; "
                                 "remove one first.")
            text = extract_text(filename, data)
            mat = Material(uuid.uuid4().hex[:12], filename, text)
            self._items[mat.id] = mat
            return mat

    def remove(self, material_id: str) -> bool:
        with self._lock:
            return self._items.pop(material_id, None) is not None

    def clear(self):
        with self._lock:
            self._items.clear()

    def summaries(self) -> list[dict]:
        with self._lock:
            return [{"id": m.id, "filename": m.filename, "chars": len(m.text)}
                    for m in self._items.values()]

    def as_docs(self) -> list[Doc]:
        with self._lock:
            return [Doc(title=m.filename, url="", text=m.text)
                    for m in self._items.values()]


STORE = MaterialStore()
